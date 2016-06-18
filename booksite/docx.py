import re
import tempfile

from docx import Document

_pattern = re.compile(r'<([a-zA-Z0-9_]+)>')


def read_docx_tags(file):
    doc = Document(file)

    tags = []
    for p in doc.paragraphs:
        matches = re.findall(_pattern, p.text)
        tags.extend(matches)

    return set(tags)


def replace_docx_tags(file, values):
    doc = Document(file)

    for tag, replace in values:
        tag = '<' + tag + '>'
        for p in doc.paragraphs:
            for r in p.runs:
                if tag in r.text:
                    r.text = r.text.replace(tag, replace)

    temp = tempfile.TemporaryFile()
    doc.save(temp)
    temp.seek(0)
    return temp
